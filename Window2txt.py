import os
from kivy.core.audio import SoundLoader
from kivy.uix.screenmanager import Screen
from tkinter import Tk, filedialog
from docx import Document
from kivy.clock import Clock

class Window2TxtFunctions():

    def open_txt_file_left_dialog(self, *args):
        root = Tk()
        root.withdraw()
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            if file_path.endswith(".docx"):
                document = Document(file_path)
                self.ids.txt1.text = "\n".join([para.text for para in document.paragraphs])
            else:
                with open(file_path, 'r+', encoding='utf-8') as file:
                    self.ids.txt1.text = file.read()

    def save_left_text_to_file(self, *args):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(self.ids.txt1.text)
                # Если файл имеет расширение .docx, то сохраняем его в формате docx
                if file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt1.text)
                    document.save(file_path)

    def open_txt_file_right_dialog(self, *args):
        root = Tk()
        root.withdraw()
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            if file_path.endswith(".docx"):
                document = Document(file_path)
                self.ids.txt2.text = "\n".join([para.text for para in document.paragraphs])
            else:
                with open(file_path, 'r+', encoding='utf-8') as file:
                    self.ids.txt2.text = file.read()

    def save_right_text_to_file(self, *args):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(self.ids.txt1.text)

                # Если файл имеет расширение .docx, то сохраняем его в формате docx
                if file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt1.text)
                    document.save(file_path)

    def find_differences(self):
        text1 = self.ids.txt1.text
        text2 = self.ids.txt2.text

        # Split the texts into lists of words
        words1 = text1.split()
        words2 = text2.split()

        # Find the differences between the two lists
        diff = list(set(words2) - set(words1))

        # Highlight the differences in the second text
        for word in diff:
            start_pos = text2.find(word)
            end_pos = start_pos + len(word)
            if start_pos != -1:
                text2 = text2[:start_pos] + f"[|]{word}[|]" + text2[end_pos:]

        self.ids.txt2.text = text2

    def remove_differences(self):
        text2 = self.ids.txt2.text
        new_text2 = text2.replace('[|]', '')
        self.ids.txt2.text = new_text2

    # метод для увеличения размера шрифта
    def increase_font_size1(self):
        self.ids.txt1.font_size += 2
        self.ids.txt2.font_size += 2

    # метод для уменьшения размера шрифта
    def decrease_font_size1(self):
        self.ids.txt1.font_size -= 2
        self.ids.txt2.font_size -= 2
class Window2txt(Window2TxtFunctions, Screen):
    def on_start(self):
        self.paused = True
        self.current_pos = 0
        self.song_time = 0
        self.update_time = None
        self.sound = None

