#TxtFunc.py
from docx import Document
from tkinter import Tk, filedialog
import keyboard
import subprocess
import os
from kivy.clock import Clock
from kivy.uix.popup import Popup
from kivy.properties import DictProperty
from kivy.uix.boxlayout import BoxLayout


class SettingsPopup(Popup):
    shortcut_labels = DictProperty({
        'Ctrl+1': '<M1>',
        'Ctrl+2': '<M2>',
        'Ctrl+3': '<M3>',
        'Ctrl+4': '<Ж1>',
        'Ctrl+5': '<Ж2>',
        'Ctrl+6': '<Ж3>',
        'Ctrl+7': '<Неразборчивая речь>',
        'Ctrl+8': '<Н>',
        'Ctrl+9': '<О>'
    })

    def save_settings(self):
        print("save_settings")
        self.dismiss()
        print(self.shortcut_labels)

class TxtFunctions():

    def add_text_to_input(self):
        print("add_text_to_input")
        a1 = self.ids.txt1.text
        cursor_pos = self.ids.txt1.cursor[0]
        line_start = a1.rfind('\n', 0, cursor_pos) + 1  # Начало текущей строки
        text_to_insert = "<M1>"

        def insert_text(dt):
            self.ids.txt1.insert_text(text_to_insert)

        Clock.schedule_once(insert_text)
    def __init__(self, shortcut_labels=None, **kwargs):
        print("init")
        super().__init__(**kwargs)
        self.secondtext = SecondText()
        self.formatted_time = "00:00:00"
        if shortcut_labels is None:
            shortcut_labels = {
                'Ctrl+1': '<M1>',
                'Ctrl+2': '<M2>',
                'Ctrl+3': '<M3>',
                'Ctrl+4': '<Ж1>',
                'Ctrl+5': '<Ж2>',
                'Ctrl+6': '<Ж3>',
                'Ctrl+7': '<Неразборчивая речь>',
                'Ctrl+8': '<Н>',
                'Ctrl+9': '<О>'
            }
        self.shortcut_labels = shortcut_labels
        keyboard.add_hotkey("Ctrl+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+1'],))
        keyboard.add_hotkey("Ctrl+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+2'],))
        keyboard.add_hotkey("Ctrl+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+3'],))
        keyboard.add_hotkey("Ctrl+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+4'],))
        keyboard.add_hotkey("Ctrl+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+5'],))
        keyboard.add_hotkey("Ctrl+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+6'],))
        keyboard.add_hotkey("Ctrl+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+7'],))
        keyboard.add_hotkey("Ctrl+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+8'],))
        keyboard.add_hotkey("Ctrl+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+9'],))
        keyboard.add_hotkey("Ctrl+0", self.print_formatted_time)


    def set_formatted_time(self, formatted_time):
        print("set_formatted_time")
        self.formatted_time = formatted_time
        print(self.formatted_time)


    def print_formatted_time(self):
        print("print_formatted_time")
        time_text = ("<" + self.formatted_time + ">")
        # self.update_shortcut_labels(self)  # Pass the 'self' instance as an argument
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(time_text), 0)
        print(self.formatted_time)



    def insert_text_from_shortcut(self, text_to_insert):
        print("insert_text_from_shortcut")
        self.update_shortcut_labels(self)  # Pass the 'self' instance as an argument
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(text_to_insert), 0)


    def open_settings_popup(self):
        print("open_settings_popup")
        popup = SettingsPopup(shortcut_labels=self.shortcut_labels)
        popup.bind(on_dismiss=self.update_shortcut_labels)  # Bind the update_shortcut_labels method to on_dismiss event
        popup.open()

    def update_shortcut_labels(self, instance):
        print("update_shortcut_labels")
        self.shortcut_labels = instance.shortcut_labels
        keyboard.remove_all_hotkeys()  # Remove all existing hotkeys

        # Add new hotkeys based on the updated shortcut_labels
        keyboard.add_hotkey("Ctrl+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+1'],))
        keyboard.add_hotkey("Ctrl+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+2'],))
        keyboard.add_hotkey("Ctrl+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+3'],))
        keyboard.add_hotkey("Ctrl+ 4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+4'],))
        keyboard.add_hotkey("Ctrl+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+5'],))
        keyboard.add_hotkey("Ctrl+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+6'],))
        keyboard.add_hotkey("Ctrl+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+7'],))
        keyboard.add_hotkey("Ctrl+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+8'],))
        keyboard.add_hotkey("Ctrl+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+9'],))
        keyboard.add_hotkey("Ctrl+0", self.print_formatted_time)
        def spech_recogn(self):
            # Получаем путь к текущей директории
            current_dir = os.path.dirname(os.path.abspath(__file__))

            # Формируем путь к папке auto_text_from_audio и к .exe файлу
            exe_dir = os.path.join(current_dir, "auto_text_from_audio")
            exe_path = os.path.join(exe_dir, "voice-recognition.exe")

            # Запускаем .exe файл в отдельном процессе
            subprocess.Popen(exe_path, creationflags=subprocess.CREATE_NEW_CONSOLE)

        def txt2(self):
            secondtext = SecondText()
            if self.secondtext in self.ids.bta.children:
                self.ids.bta.remove_widget(self.secondtext)
            else:
                self.ids.bta.add_widget(self.secondtext)

        def open_txt_file_dialog(self, *args):
            root = Tk()
            root.withdraw()
            file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
            if file_path:
                if file_path.endswith(".docx"):
                    document = Document(file_path)
                    self.ids.txt1.text = "\n".join([para.text for para in document.paragraphs])
                else:
                    with open(file_path, 'r+', encoding='cp1251') as file:
                        self.ids.txt1.text = file.read()

        def save_text_to_file(self, *args):
            file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(self.ids.txt1.text)
                    # Если файл имеет расширение .docx, то сохраняем его в формате docx
                    if file_path.endswith(".docx"):
                        document = Document()
                        document.add_paragraph(self.ids.txt1.text)
                        document.save(file_path)

        # метод для увеличения размера шрифта
        def increase_font_size1(self):
            self.ids.txt1.font_size += 2

        # метод для уменьшения размера шрифта
        def decrease_font_size1(self):
            self.ids.txt1.font_size -= 2

class SecondText(BoxLayout):
    def open_22_file_dialog(self, *args):
        root = Tk()
        root.withdraw()
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            if file_path.endswith(".docx"):
                document = Document(file_path)
                self.ids.txt22.text = "\n".join([para.text for para in document.paragraphs])
            else:
                with open(file_path, 'r+', encoding='utf-8') as file:
                    self.ids.txt22.text = file.read()

    def save_22_text_to_file(self, *args):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(self.ids.txt22.text)
                # Если файл имеет расширение .docx, то сохраняем его в формате docx
                if file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt22.text)
                    document.save(file_path)

    def increase_font_size2(self):
        self.ids.txt22.font_size += 2

    # метод для уменьшения размера шрифта
    def decrease_font_size2(self):
        self.ids.txt22.font_size -= 2