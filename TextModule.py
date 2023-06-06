# TextModule.py
from docx import Document
from tkinter import Tk, filedialog
import keyboard
import subprocess
import os
import time
from kivy.clock import Clock
from kivy.uix.boxlayout import BoxLayout
from zipfile import BadZipFile
import win32com.client as win32

#основной класс для работы с текстом
class Text1Func():
    def __init__(self, shortcut_labels=None, **kwargs):
        print("init")
        super().__init__(**kwargs)
        self.secondtext = Text2Func()
        self.formatted_time = "00:00:00"
        if shortcut_labels is None:
            shortcut_labels = {
                'Alt+`': 'М? — ',
                'Alt+1': 'M1 — ',
                'Alt+2': 'M2 — ',
                'Alt+3': 'M3 — ',
                'Alt+4': 'M4 — ',
                'Alt+5': '<Речь неразборчива>',
                'Alt+6': 'Н — ',
                'Alt+7': 'О — ',
                'Alt+8': '<>',
                'Alt+9': 'Спорная фонограмма №',

                'Ctrl+`': 'Ж? — ',
                'Ctrl+1': 'Ж1 — ',
                'Ctrl+2': 'Ж2 — ',
                'Ctrl+3': 'Ж3 — ',
                'Ctrl+4': 'Ж4 — ',
                'Ctrl+5': 'Р — ',
                'Ctrl+6': '<>',
                'Ctrl+7': '<>',
                'Ctrl+8': '<>',
                'Ctrl+9': '<>',
                'Ctrl+0': '<Комментарий эксперта>'
            }
        self.shortcut_labels = shortcut_labels
        keyboard.add_hotkey("Ctrl+`", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+`'],))
        keyboard.add_hotkey("Ctrl+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+1'],))
        keyboard.add_hotkey("Ctrl+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+2'],))
        keyboard.add_hotkey("Ctrl+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+3'],))
        keyboard.add_hotkey("Ctrl+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+4'],))
        keyboard.add_hotkey("Ctrl+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+5'],))
        keyboard.add_hotkey("Ctrl+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+6'],))
        keyboard.add_hotkey("Ctrl+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+7'],))
        keyboard.add_hotkey("Ctrl+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+8'],))
        keyboard.add_hotkey("Ctrl+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+9'],))
        keyboard.add_hotkey("Ctrl+0", self.insert_text_from_shortcut, args=(self.shortcut_labels['Ctrl+0'],))
        keyboard.add_hotkey("Alt+`", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+`'],))
        keyboard.add_hotkey("Alt+1", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+1'],))
        keyboard.add_hotkey("Alt+2", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+2'],))
        keyboard.add_hotkey("Alt+3", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+3'],))
        keyboard.add_hotkey("Alt+4", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+4'],))
        keyboard.add_hotkey("Alt+5", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+5'],))
        keyboard.add_hotkey("Alt+6", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+6'],))
        keyboard.add_hotkey("Alt+7", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+7'],))
        keyboard.add_hotkey("Alt+8", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+8'],))
        keyboard.add_hotkey("Alt+9", self.insert_text_from_shortcut, args=(self.shortcut_labels['Alt+9'],))
        keyboard.add_hotkey("Alt+0", self.print_formatted_time)

    #вствка текста зависящего от нажатой кнопки в основной текст
    def add_text_to_input(self):
        a1 = self.ids.txt1.text
        cursor_pos = self.ids.txt1.cursor[0]
        line_start = a1.rfind('\n', 0, cursor_pos) + 1
        text_to_insert = ""

        def insert_text(dt):
            self.ids.txt1.insert_text(text_to_insert)
        Clock.schedule_once(insert_text)

    #метод для обновления текушего времени воспроизведения
    def set_formatted_time(self, formatted_time):
        self.formatted_time = formatted_time

    # метод для вставки текушего времени воспроизведения в качестве текстовой метки
    def print_formatted_time(self):
        time_text = ("<" + self.formatted_time + ">")
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(time_text), 0)
    #метод для вставки текстовых меток

    def insert_text_from_shortcut(self, text_to_insert):
        Clock.schedule_once(lambda dt: self.ids.txt1.insert_text(text_to_insert), 0)

    #метод для перревода интерфейса в режим работы с двумя текстами
    def txt2(self):
        secondtext = Text2Func()
        if self.secondtext in self.ids.bta.children:
            self.ids.bta.remove_widget(self.secondtext)
        else:
            self.ids.bta.add_widget(self.secondtext)

    #открытие текстовых файлов через диалоговое окно
    def open_txt_file_dialog(self, *args):
        try:
            root = Tk()
            root.withdraw()
            file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"),
                                                              ("Word Documents", "*.docx"),
                                                              ("Word Documents (Legacy)", "*.doc")])
            if file_path:
                if file_path.endswith(".docx"):
                    document = Document(file_path)
                    paragraphs = [para.text for para in document.paragraphs]
                    text = "\n".join(paragraphs)
                    self.ids.txt1.text = text
                elif file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(file_path)
                    text = doc.Content.Text
                    self.ids.txt1.text = text
                    doc.Close()
                    word.Quit()
                else:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                        self.ids.txt1.text = text
        except BadZipFile:
            pass
        except FileNotFoundError:
            pass
        except IOError:
            pass
        except PermissionError:
            pass
        except Exception as e:
            pass

    #сохранение текстовых файлов через диалоговое окно
    def save_text_to_file(self, *args):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text Files", "*.txt"),
                                                                ("Word Documents", "*.docx"),
                                                                ("Word Documents (Legacy)", "*.doc")])
            if file_path:
                if file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt1.text)
                    document.save(file_path)
                elif file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Add()
                    doc.Content.Text = self.ids.txt1.text
                    doc.SaveAs(file_path)
                    doc.Close()
                    word.Quit()
                else:
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(self.ids.txt1.text)
        except BadZipFile:
            pass
        except FileNotFoundError:
            pass
        except IOError:
            pass
        except PermissionError:
            pass
        except Exception as e:
            pass

    #увеличение размера шрифта первого текста
    def increase_font_size1(self):
        self.ids.txt1.font_size += 2

    #уменьшение размера шрифта первого текста
    def decrease_font_size1(self):
        self.ids.txt1.font_size -= 2

    # запуск модуля распознования речи
    def spech_recogn(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        exe_dir = os.path.join(current_dir, "auto_text_from_audio")
        exe_path = os.path.join(exe_dir, "voice-recognition.exe")
        files_dir = os.path.join(exe_dir, "vosk-model-ru-0.42")
        subprocess.Popen([exe_path, files_dir], creationflags=subprocess.CREATE_NEW_CONSOLE, cwd=exe_dir)

#класс для работы с вторым текстом
class Text2Func(BoxLayout):
    #открытие текстового файла второго текста
    def open_22_file_dialog(self, *args):
        try:
            root = Tk()
            root.withdraw()
            file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"),
                                                              ("Word Documents", "*.docx"),
                                                              ("Word Documents (Legacy)", "*.doc")])
            if file_path:
                if file_path.endswith(".docx"):
                    document = Document(file_path)
                    paragraphs = [para.text for para in document.paragraphs]
                    text = "\n".join(paragraphs)
                    self.ids.txt22.text = text
                elif file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(file_path)
                    text = doc.Content.Text
                    self.ids.txt22.text = text
                    doc.Close()
                    word.Quit()
                else:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                        self.ids.txt22.text = text
        except BadZipFile:
            pass
        except FileNotFoundError:
            pass
        except IOError:
            pass
        except PermissionError:
            pass
        except Exception as e:
            pass

    # сохранение текстового файла для второго текста
    def save_22_text_to_file(self, *args):
        try:
            file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                     filetypes=[("Text Files", "*.txt"),
                                                                ("Word Documents", "*.docx"),
                                                                ("Word Documents (Legacy)", "*.doc")])
            if file_path:
                if file_path.endswith(".docx"):
                    document = Document()
                    document.add_paragraph(self.ids.txt22.text)
                    document.save(file_path)
                elif file_path.endswith(".doc"):
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Add()
                    doc.Content.Text = self.ids.txt22.text
                    doc.SaveAs(file_path)
                    doc.Close()
                    word.Quit()
                else:
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(self.ids.txt22.text)
        except BadZipFile:
            pass
        except FileNotFoundError:
            pass
        except IOError:
            pass
        except PermissionError:
            pass
        except Exception as e:
            pass
    #увеличение размера шрифта второго текста
    def increase_font_size2(self):
        self.ids.txt22.font_size += 2

    #уменьшение размера шрифта второго текста
    def decrease_font_size2(self):
        self.ids.txt22.font_size -= 2
